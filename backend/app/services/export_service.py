"""Exportación de traducciones a TXT, DOCX y PDF."""
from __future__ import annotations

import io
from datetime import datetime
from typing import Sequence

from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from app.models.translation import Translation


def _fmt(ts: datetime) -> str:
    return ts.strftime("%Y-%m-%d %H:%M")


def export_txt(translations: Sequence[Translation]) -> bytes:
    lines = ["SISTEMA LSC i5.0 — EXPORTACIÓN DE TRADUCCIONES", "=" * 50, ""]
    for t in translations:
        lines.append(f"[{_fmt(t.created_at)}] (conf={t.confidence:.2f})")
        lines.append(f"  Glosas : {' '.join(t.glosses or [])}")
        lines.append(f"  Texto  : {t.natural_text}")
        lines.append("")
    return "\n".join(lines).encode("utf-8")


def export_docx(translations: Sequence[Translation]) -> bytes:
    doc = Document()
    title = doc.add_heading("Sistema LSC i5.0 — Traducciones", level=0)
    title.runs[0].font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

    doc.add_paragraph(f"Generado: {_fmt(datetime.utcnow())}")
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Fecha", "Glosas", "Texto natural", "Conf."]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].font.bold = True

    for t in translations:
        row = table.add_row().cells
        row[0].text = _fmt(t.created_at)
        row[1].text = " ".join(t.glosses or [])
        row[2].text = t.natural_text
        row[3].text = f"{t.confidence:.2f}"

    for p in doc.paragraphs:
        for run in p.runs:
            run.font.size = run.font.size or Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_pdf(translations: Sequence[Translation]) -> bytes:
    buf = io.BytesIO()
    docp = SimpleDocTemplate(buf, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm)
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="LSCTitle", parent=styles["Title"],
            textColor=colors.HexColor("#4F46E5"),
        )
    )
    story = [
        Paragraph("Sistema LSC i5.0 — Traducciones", styles["LSCTitle"]),
        Spacer(1, 8),
        Paragraph(f"Generado: {_fmt(datetime.utcnow())}", styles["Normal"]),
        Spacer(1, 16),
    ]

    data = [["Fecha", "Glosas", "Texto natural", "Conf."]]
    for t in translations:
        data.append([
            _fmt(t.created_at),
            " ".join(t.glosses or []),
            Paragraph(t.natural_text, styles["Normal"]),
            f"{t.confidence:.2f}",
        ])

    table = Table(data, colWidths=[3 * cm, 4 * cm, 7 * cm, 1.8 * cm], repeatRows=1)
    table.setStyle(
        TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4F46E5")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CBD5E1")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F1F5F9")]),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ])
    )
    story.append(table)
    docp.build(story)
    return buf.getvalue()


EXPORTERS = {
    "txt": (export_txt, "text/plain", "txt"),
    "docx": (
        export_docx,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "docx",
    ),
    "pdf": (export_pdf, "application/pdf", "pdf"),
}
